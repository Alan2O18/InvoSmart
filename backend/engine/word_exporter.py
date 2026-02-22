import copy
import logging
from typing import Dict, Any, List
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from backend.repositories.project_repository import ProjectRepository
from backend.repositories.job_repository import JobRepository

logger = logging.getLogger(__name__)

class WordExporter:
    """Word 報表匯出模組 (依品類跨頁分組)"""

    def __init__(self, project_repo: ProjectRepository):
        self.project_repo = project_repo

    def _replace_text_in_paragraph(self, paragraph, replacements: Dict[str, str], mark_unfilled_red: bool = True):
        """
        處理 Word 中的段落替換，為避免 Run 被打斷（如 {{ 和 }} 分在不同 Run），
        會將整個 Paragraph 的文字先提取處理後再回填，
        若保留的 {{}} 佔位符需要標記，則標為紅色。
        """
        # 因為 p.text 的替換會丟失原有的樣式，較安全的做法是合併 Runs 或操作文本。
        # 最簡單暴力的完整保留樣式法：檢查 text 是否有大括號
        if "{{" not in paragraph.text and "活動總務：李天旭" not in paragraph.text:
            return

        text = paragraph.text
        # 取代硬編碼
        if "活動總務：李天旭" in text:
            coordinator = replacements.get("{{活動總務}}") or replacements.get("{{活動總召}}", "")
            if coordinator:
                text = text.replace("活動總務：李天旭", f"活動總務：{coordinator}")
            elif mark_unfilled_red:
                text = text.replace("活動總務：李天旭", "{{活動總務}}") # 讓它變成佔位符以等一下染成紅色
            else:
                text = text.replace("活動總務：李天旭", "")
        
        # 進行常規變數替換
        for key, val in replacements.items():
            if key in text:
                if val:
                    text = text.replace(key, str(val))
                elif not mark_unfilled_red:
                    text = text.replace(key, "")

        # 寫回 paragraph
        # 由於直接改 paragraph.text 會清除所有 runs 的樣式，
        # 這裡採用替換首個 run 的內容，並清除其餘 runs，這是常見的 workaround。
        if len(paragraph.runs) > 0:
            paragraph.runs[0].text = text
            for idx in range(1, len(paragraph.runs)):
                paragraph.runs[idx].text = ""
            
            # 若 mark_unfilled_red 打開，並且還是找到了 {{}}
            # 在 python-docx 裡，如果要對特定字眼塗紅不破壞其他字，
            # 需要把 paragraph 字串再拆分。
            if mark_unfilled_red and "{{" in text:
                self._highlight_unfilled_placeholders(paragraph)

    def _highlight_unfilled_placeholders(self, paragraph):
        """將段落中的 {{...}} 標為紅色"""
        text = paragraph.text
        if "{{" not in text:
            return
            
        # 清除所有 runs
        for run in paragraph.runs:
            run.text = ""
            
        import re
        # 分割字串，保留 {{...}} 作為獨立片段
        parts = re.split(r'(\{\{.+?\}\})', text)
        
        for part in parts:
            if not part: continue
            run = paragraph.add_run(part)
            if part.startswith("{{") and part.endswith("}}"):
                run.font.color.rgb = RGBColor(255, 0, 0) # 紅色

    def _replace_text_in_table(self, table, replacements: Dict[str, str], mark_unfilled_red: bool = True):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    self._replace_text_in_paragraph(p, replacements, mark_unfilled_red)
        
    def _find_row_with_placeholder(self, table, placeholder) -> int:
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                if placeholder in cell.text:
                    return i
        return -1

    def _copy_table_after(self, table, paragraph):
        """複製一個表格，包含其所有 XML 結構 (保留合併儲存格等複雜樣式)，並插入到指定的 paragraph 後面"""
        tbl, p = table._tbl, paragraph._p
        new_tbl = copy.deepcopy(tbl)
        p.addnext(new_tbl)
        return new_tbl

    def _add_page_break(self, doc):
        doc.add_page_break()

    def process_export(self, project_id: str, template_path: str, job_repo) -> str:
        # 1. 取得專案元資料
        project = self.project_repo.get_project(project_id)
        if not project:
            raise ValueError(f"Project not found: {project_id}")
            
        meta = project.get("metadata", {})
        
        # 2. 獲取所有 Job，扁平化 Items 並分類
        jobs = job_repo.list_jobs()
        
        grouped_items = {}
        import json
        for job in jobs:
            vlm_raw = job.get("vlm_result_json")
            vlm = {}
            if isinstance(vlm_raw, str) and vlm_raw.strip():
                try:
                    vlm = json.loads(vlm_raw)
                except:
                    pass
            elif isinstance(vlm_raw, dict):
                vlm = vlm_raw
                
            items = vlm.get("items", []) if isinstance(vlm, dict) else []
            header = vlm.get("header", {}) if isinstance(vlm, dict) else {}
            
            # 兼容舊版可能不是 dict 的狀況
            if not isinstance(header, dict):
                header = {}
            if not isinstance(items, list):
                items = []
                
            voucher_id = header.get("voucher_id", "") if isinstance(header, dict) else ""
            
            for item in items:
                cat = item.get("category", "未分類") or "未分類"
                if cat not in grouped_items:
                    grouped_items[cat] = []
                    
                flat_item = {
                    "voucher_id": voucher_id,
                    "name": item.get("name") or item.get("description", ""),
                    "qty": item.get("qty") or item.get("quantity", ""),
                    "price": item.get("price", ""),
                    "total": item.get("total", ""),
                    "purpose": ""  # 預留用途
                }
                grouped_items[cat].append(flat_item)
                
        # 3. 操作 Word
        doc = Document(template_path)
        
        # 運用 XML 節點 deepcopy 達成完整頁面的複製 (針對範本文件僅有單頁之情形)
        body = doc.element.body
        sect_pr = body.xpath('./w:sectPr')
        sect_pr = sect_pr[0] if sect_pr else None
        
        template_elements = []
        for e in list(body.iterchildren()):
            if e.tag.endswith('sectPr'):
                continue
            template_elements.append(copy.deepcopy(e))
            body.remove(e)

        categories = list(grouped_items.keys())
        if not categories:
            categories = ["無報帳記錄"]
            
        # 準備通用的替換字典 (Metadata)
        t_count = int(meta.get("teacherCount", 0) or 0)
        s_count = int(meta.get("studentCount", 0) or 0)
        total_count = t_count + s_count
        
        start_time = meta.get("startTime", "")
        end_time = meta.get("endTime", "")
        period = f"{start_time} ~ {end_time}" if (start_time and end_time) else (start_time or "")
            
        base_replacements = {
            "{{組別}}": meta.get("group", ""),
            "{{組長}}": meta.get("leader", ""),
            "{{活動名稱}}": meta.get("name", ""),
            "{{活動總召}}": meta.get("coordinator", ""),
            "{{活動總務}}": meta.get("generalAffairs", ""),
            "{{活動期間}}": period,
            "{{活動地點}}": meta.get("location", ""),
            "{{總人數}}": str(total_count) if total_count else "",
            "{{老師人數}}": str(t_count) if t_count else "",
            "{{學生人數}}": str(s_count) if s_count else "",
        }

        for idx, cat in enumerate(categories):
            old_p_len = len(doc.paragraphs)
            
            for e in template_elements:
                new_e = copy.deepcopy(e)
                if sect_pr is not None:
                    sect_pr.addprevious(new_e)
                else:
                    body.append(new_e)

            # Metadata for current page
            rep = dict(base_replacements)
            if len(categories) > 1 and cat != "無報帳記錄":
                rep["{{活動名稱}}"] = f"{meta.get('name', '')} - {cat}"
            else:
                rep["{{活動名稱}}"] = meta.get('name', '')
                
            # 處理剛插入頁面的表格 (Table 0 預算表, Table 1 結算表)
            t0 = doc.tables[idx * 2] if (idx * 2) < len(doc.tables) else None
            t1 = doc.tables[idx * 2 + 1] if (idx * 2 + 1) < len(doc.tables) else None
            
            if t0:
                self._replace_text_in_table(t0, rep)
            
            sum_total = 0
            if t1:
                # 填入明細到 t1
                row_idx = self._find_row_with_placeholder(t1, "{{決算支出列}}")
                if row_idx != -1:
                    target_row = t1.rows[row_idx]
                    target_tr = target_row._tr
                    cat_items = grouped_items.get(cat, [])
                    
                    for item in cat_items:
                        new_tr = copy.deepcopy(target_tr)
                        target_tr.addprevious(new_tr)
                        
                        from docx.table import _Row
                        new_row = _Row(new_tr, t1)
                        
                        try:
                            val = item.get("total", 0)
                            if val == "": val = 0
                            t = float(val)
                            sum_total += int(t)
                        except Exception:
                            pass
                            
                        cells = new_row.cells
                        if len(cells) >= 6:
                            cells[0].text = str(item["name"])
                            cells[1].text = str(item["qty"]) if item["qty"] else ""
                            cells[2].text = str(item["price"]) if item["price"] else ""
                            cells[3].text = str(item["total"]) if item["total"] else ""
                            cells[4].text = str(item["purpose"])
                            cells[5].text = str(item["voucher_id"])
                            
                    # 移除模板的那一行 {{決算支出列}}
                    target_tr.getparent().remove(target_tr)
                    
                rep["{{決算_支出總計}}"] = str(sum_total)
                self._replace_text_in_table(t1, rep)
            
            # 替換此頁段落的變數 (如標題)
            new_p_len = len(doc.paragraphs)
            for i in range(old_p_len, new_p_len):
                self._replace_text_in_paragraph(doc.paragraphs[i], rep)
                
            # 若不是最後一頁，強制插入分頁符號
            if idx < len(categories) - 1:
                from docx.oxml import OxmlElement
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                br = OxmlElement('w:br')
                br.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'page')
                r.append(br)
                p.append(r)
                if sect_pr is not None:
                    sect_pr.addprevious(p)
                else:
                    body.append(p)
                    
        # 4. 儲存實體檔案
        out_root = self.project_repo._project_root(project_id) / "Word匯出"
        out_root.mkdir(parents=True, exist_ok=True)
        fname = f"{project_id}_word_export.docx"
        out_path = out_root / fname
        doc.save(str(out_path))
        
        return str(out_path)
