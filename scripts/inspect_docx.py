from docx import Document

try:
    doc = Document(r"backend/data/workspace/T0/Word匯出/T0_word_export.docx")
    with open("inspect_docx_out_utf8.txt", "w", encoding="utf-8") as f:
        for i, t in enumerate(doc.tables):
            f.write(f"\n--- Table {i} ---\n")
            for r, row in enumerate(t.rows):
                cell_texts = [c.text.replace('\n', ' ').strip() for c in row.cells]
                f.write(f"Row {r}: {cell_texts}\n")

        f.write("\n--- Paragraphs ---\n")
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                f.write(f"Para {i}: {p.text.strip()}\n")
except Exception as e:
    with open("inspect_docx_out_utf8.txt", "w", encoding="utf-8") as f:
        f.write(f"Error: {e}\n")
