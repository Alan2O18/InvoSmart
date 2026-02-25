import os
from docx import Document
from backend.engine.core import Engine

def test_word_export():
    eng = Engine(start_workers=False)
    # Check if there's any completed project
    projects = eng.project_repo.list_projects()
    if not projects:
        print("No projects to test.")
        return

    # Let's take the first project to test
    p_id = projects[0]["project_id"]
    print(f"Testing export for project: {p_id}")
    
    template_path = "dev_data/空白 模板 (1).docx"
    try:
        out_path = eng.export_handler.run_word(p_id, template_path)
        print("Successfully generated Word file:", out_path)
        
        # Verify document looks right
        doc = Document(out_path)
        print(f"Generated doc has {len(doc.tables)} tables and {len(doc.paragraphs)} paragraphs.")
    except Exception as e:
        print("Failed to export:", e)

if __name__ == "__main__":
    test_word_export()
