import re
from docx import Document
import io

doc = Document('dev_data/空白 模板 (1).docx')

texts = []
for p in doc.paragraphs:
    texts.append(p.text)
    
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            texts.append(c.text)

all_text = " ".join(texts)
placeholders = set(re.findall(r'\{\{.+?\}\}', all_text))

with open('check_out_utf8.txt', 'w', encoding='utf-8') as f:
    f.write("Placeholders found: " + str(placeholders) + "\n\n")
    f.write("--- Table 0 ---\n")
    t0 = doc.tables[0]
    for r in t0.rows:
        f.write(str([c.text.strip() for c in r.cells]) + "\n")
