import os
import re
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document
from docx.shared import RGBColor

from backend.engine.word_exporter import WordExporter


@pytest.fixture
def word_exporter():
    project_repo_mock = MagicMock()
    return WordExporter(project_repo_mock)


def test_replace_text_in_paragraph(word_exporter):
    doc = Document()
    paragraph = doc.add_paragraph("This is a {{test_key}} paragraph.")

    replacements = {"{{test_key}}": "successful"}
    word_exporter._replace_text_in_paragraph(paragraph, replacements, mark_unfilled_red=False)

    assert "successful" in paragraph.text
    assert "{{test_key}}" not in paragraph.text


def test_replace_text_hardcoded_coordinator(word_exporter):
    doc = Document()
    paragraph = doc.add_paragraph("負責人是 活動總務：李天旭 先生")

    replacements = {"{{活動總務}}": "王小明"}
    word_exporter._replace_text_in_paragraph(paragraph, replacements, mark_unfilled_red=False)

    assert "活動總務：王小明" in paragraph.text
    assert "李天旭" not in paragraph.text


def test_highlight_unfilled_placeholders(word_exporter):
    doc = Document()
    paragraph = doc.add_paragraph("Missing {{data_key}} here.")

    word_exporter._replace_text_in_paragraph(paragraph, {}, mark_unfilled_red=True)

    assert len(paragraph.runs) > 1
    assert any(
        "{{data_key}}" in run.text and run.font.color and run.font.color.rgb == RGBColor(255, 0, 0)
        for run in paragraph.runs
    )


def test_replace_text_in_table(word_exporter):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Data: {{val}}"

    word_exporter._replace_text_in_table(table, {"{{val}}": "100"})
    assert "100" in table.cell(0, 0).text


def test_find_row_with_placeholder(word_exporter):
    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Row 1"
    table.cell(1, 0).text = "{{TARGET}}"

    assert word_exporter._find_row_with_placeholder(table, "{{TARGET}}") == 1
    assert word_exporter._find_row_with_placeholder(table, "{{MISSING}}") == -1


def test_set_cell_text_formatted(word_exporter):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    word_exporter._set_cell_text_formatted(cell, "Formatted Text")
    assert cell.text == "Formatted Text"


def test_format_roc_date(word_exporter):
    assert word_exporter._format_roc_date("2024-05-12") == "113年5月12日"
    assert word_exporter._format_roc_date("invalid") == "invalid"


def test_normalize_people_names_supports_multi_formats(word_exporter):
    assert word_exporter._normalize_people_names("王大明、李小華") == "王大明、李小華"
    assert word_exporter._normalize_people_names("王大明, 李小華") == "王大明、李小華"
    assert word_exporter._normalize_people_names('["王大明", "李小華", "王大明"]') == "王大明、李小華"
    assert word_exporter._normalize_people_names(["王大明", "李小華", "王大明"]) == "王大明、李小華"


def test_format_activity_period_uses_roc_lines(word_exporter):
    period = word_exporter._format_activity_period("2026-03-20T09:30", "2026-03-20T12:00")
    assert period == "自民國115年 3月 20日 09:30分(開始)\n到民國115年 3月 20日 12:00分(結束)"


@pytest.mark.asyncio
async def test_process_export_no_project_raises(word_exporter):
    word_exporter.project_repo = AsyncMock()
    word_exporter.project_repo.get_project.return_value = None

    with pytest.raises(ValueError, match="Project not found: p1"):
        await word_exporter.process_export("p1", "dummy.docx", AsyncMock())


@pytest.mark.asyncio
async def test_ensure_flatten_cache_uses_display_result_and_category_priority(tmp_path):
    project_repo = AsyncMock()
    project_repo.get_project.return_value = {"metadata": {"group": "教材費"}}
    project_repo._project_root = MagicMock(return_value=tmp_path)

    exporter = WordExporter(project_repo)

    job_repo = AsyncMock()
    job_repo.project_id = "proj_sort"
    job_repo.list_jobs.return_value = [{"job_id": "job1", "updated_at": 20}]
    job_repo.get_display_result = AsyncMock(
        return_value={
            "header": {"voucher_id": "V1", "supplier": "供應商A", "date": "2026-04-01"},
            "summary": {"purpose": "用途A", "total": 100},
            "items": [
                {"category": "其他雜項", "name": "項目B", "qty": 1, "price": 60, "total": 60},
                {"category": "保險", "name": "項目A", "qty": 1, "price": 40, "total": 40},
            ],
        }
    )

    payload = await exporter.ensure_flatten_cache("proj_sort", job_repo)

    assert payload["sumTotal"] == 100
    assert payload["payloadSources"]["display_result"] == 1
    assert payload["allFlattenedItems"][0]["_category"] == "保險"


@pytest.mark.asyncio
async def test_process_export_filename_and_activity_name_format(word_exporter, tmp_path):
    job_repo = AsyncMock()
    job_repo.project_id = "proj_123"
    job_repo.list_jobs.return_value = [{"job_id": "job1", "updated_at": 100}]
    job_repo.get_display_result = AsyncMock(
        return_value={
            "header": {"voucher_id": "VX-01", "supplier": "店家", "date": "2026-04-01"},
            "summary": {"purpose": "活動用品", "total": 200},
            "items": [
                {"category": "餐食", "name": "餐盒", "qty": 2, "price": 100, "total": 200},
            ],
        }
    )

    word_exporter.project_repo = AsyncMock()
    word_exporter.project_repo.get_project.return_value = {
        "metadata": {
            "name": "春季/成果:發表?*",
            "coordinator": "測試總召",
            "budgetIncome": [{"amount": "500", "name": "補助"}],
            "budgetExpense": [{"total": "300", "name": "餐費"}],
        },
        "root_path": str(tmp_path),
    }
    word_exporter.project_repo._project_root = MagicMock(return_value=tmp_path)

    template_path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("活動：{{活動名稱}}")
    # Keep two tables to cover replacement branches.
    table1 = doc.add_table(rows=1, cols=5)
    table1.cell(0, 0).text = "{{預算支出列}}"
    table2 = doc.add_table(rows=1, cols=6)
    table2.cell(0, 0).text = "{{決算支出列}}"
    doc.save(str(template_path))

    out_path = await word_exporter.process_export("proj_123", str(template_path), job_repo)

    assert os.path.exists(out_path)
    out_name = os.path.basename(out_path)
    assert re.match(r"^proj_123「春季_成果_發表__」_預結算表_\d{8}_\d{6}\.docx$", out_name)

    exported_doc = Document(out_path)
    full_text = "\n".join(paragraph.text for paragraph in exported_doc.paragraphs)
    assert "活動：proj_123 春季/成果:發表?*" in full_text
